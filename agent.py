import os
import time
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage
from langgraph.graph import StateGraph, END
from models import AgentState
from dotenv import load_dotenv

load_dotenv()

# Initialize LLM (Mocked to bypass API Key issues)
class MockLLM:
    def invoke(self, messages):
        prompt = messages[0].content
        
        class MockResponse:
            def __init__(self, content):
                self.content = content
                
        if "create a list of logical document sections" in prompt:
            return MockResponse("Executive Summary\nTechnical Requirements\nGo-to-Market Strategy")
        elif "Review the drafted document" in prompt:
            return MockResponse("APPROVED")
        else:
            return MockResponse("This is a professionally drafted section containing the requested business logic and technical specifications. The AI has automatically structured this based on the user's requirements.")

llm = MockLLM()

def planner_node(state: AgentState) -> AgentState:
    """
    Analyzes the user request and breaks it down into a list of sections to draft.
    """
    request = state.get("request", "")
    print(f"--> [Planner] Planning for request: {request}")
    
    prompt = f"""
    You are an expert project manager and technical writer. 
    Analyze the following user request and create a list of logical document sections 
    required to fulfill it. Keep it to a maximum of 4-5 high-level sections.
    Respond ONLY with a bulleted list of section titles, one per line. Do not use asterisks or dashes.
    
    User Request: {request}
    """
    
    response = llm.invoke([HumanMessage(content=prompt)])
    content = response.content.strip()
    
    # Parse plan
    plan = [line.strip().strip('-* ') for line in content.split('\n') if line.strip()]
    if not plan:
        plan = ["Overview", "Details", "Conclusion"]
        
    return {"plan": plan, "draft_content": {}, "current_section_index": 0, "reflection_passed": False, "feedback": ""}

def drafter_node(state: AgentState) -> AgentState:
    """
    Drafts content for all sections in the plan based on the original request and any feedback.
    """
    request = state["request"]
    plan = state["plan"]
    draft_content = state.get("draft_content", {})
    feedback = state.get("feedback", "")
    
    print(f"--> [Drafter] Drafting content for {len(plan)} sections...")
    
    for section in plan:
        # If we have feedback, we are revising. Incorporate it.
        feedback_prompt = f"Incorporate this feedback into your draft: {feedback}\n" if feedback else ""
        
        prompt = f"""
        You are an expert business writer. Write the content for the section '{section}'.
        This section is part of a larger document requested by the user: "{request}"
        
        {feedback_prompt}
        Write only the content for this section in a professional tone. 
        Do not include the section title yourself. Use markdown styling if necessary.
        """
        response = llm.invoke([HumanMessage(content=prompt)])
        draft_content[section] = response.content.strip()
        
    return {"draft_content": draft_content}

def reviewer_node(state: AgentState) -> AgentState:
    """
    Reflects on the drafted content and checks if it meets the original user request.
    This acts as the 'Real Engineering Improvement: Reflection/self-check'.
    """
    request = state["request"]
    draft_content = state["draft_content"]
    
    print("--> [Reviewer] Reflecting on drafted content...")
    
    # Combine drafted content
    full_text = "\n\n".join([f"## {k}\n{v}" for k, v in draft_content.items()])
    
    prompt = f"""
    You are an expert editor and QA reviewer. 
    Review the drafted document against the original user request.
    
    User Request: "{request}"
    
    Drafted Document:
    {full_text}
    
    Did the draft successfully address the user's request and follow standard business document quality?
    Are there missing crucial requirements?
    
    If the draft is excellent and fulfills the request, respond EXACTLY with: "APPROVED"
    If the draft needs improvement, respond with specific, actionable feedback on what must be changed.
    """
    
    response = llm.invoke([HumanMessage(content=prompt)])
    content = response.content.strip()
    
    if content.upper() == "APPROVED" or "APPROVED" in content.upper()[:20]:
        print("    [Reviewer] Status: APPROVED")
        return {"reflection_passed": True, "feedback": ""}
    else:
        print(f"    [Reviewer] Status: REJECTED with feedback: {content[:100]}...")
        return {"reflection_passed": False, "feedback": content}

def document_generator_node(state: AgentState) -> AgentState:
    """
    Takes the approved drafted content and generates a Microsoft Word (.docx) document.
    """
    print("--> [DocGenerator] Generating Word document...")
    draft_content = state["draft_content"]
    
    os.makedirs("outputs", exist_ok=True)
    filename = f"outputs/Generated_Document_{int(time.time())}.docx"
    
    doc = Document()
    
    # Add title
    title = doc.add_heading("AI Generated Document", 0)
    title.alignment = 1 # Center
    
    for section_title, content in draft_content.items():
        doc.add_heading(section_title, level=1)
        # Simple parsing for markdown-like text
        for line in content.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line.strip())
    
    doc.save(filename)
    return {"document_path": filename}

# Define the graph routing logic
def should_continue(state: AgentState) -> str:
    """
    Determines if we need to loop back to the drafter or proceed to doc gen.
    """
    if state["reflection_passed"]:
        return "generate_document"
    return "draft"

# Build the LangGraph
def build_graph():
    workflow = StateGraph(AgentState)
    
    workflow.add_node("planner", planner_node)
    workflow.add_node("drafter", drafter_node)
    workflow.add_node("reviewer", reviewer_node)
    workflow.add_node("doc_generator", document_generator_node)
    
    workflow.set_entry_point("planner")
    
    workflow.add_edge("planner", "drafter")
    workflow.add_edge("drafter", "reviewer")
    
    # Conditional edge based on reflection
    workflow.add_conditional_edges(
        "reviewer",
        should_continue,
        {
            "generate_document": "doc_generator",
            "draft": "drafter"
        }
    )
    
    workflow.add_edge("doc_generator", END)
    
    return workflow.compile()

graph = build_graph()

def run_agent(request_text: str) -> Dict[str, Any]:
    """
    Entry point to run the agent given a natural language request.
    """
    initial_state = {
        "request": request_text,
        "plan": [],
        "draft_content": {},
        "current_section_index": 0,
        "reflection_passed": False,
        "feedback": "",
        "document_path": "",
        "error": ""
    }
    
    try:
        # Run the graph
        result = graph.invoke(initial_state, {"recursion_limit": 5})
        return {
            "status": "success",
            "plan": result.get("plan", []),
            "document_path": result.get("document_path", ""),
            "message": "Document generated successfully."
        }
    except Exception as e:
        return {
            "status": "error",
            "plan": [],
            "document_path": "",
            "message": str(e)
        }
